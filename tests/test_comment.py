"""Unit tests for the Comment class."""

from docx import Document


class TestComment:
    def it_works(self):
        doc = Document()
        p = doc.add_paragraph("Hello world!")
        p.add_run(" Second run!")

        # Add a comment to the paragraph's first run
        r1 = p.runs[0]
        c = r1.add_comment("New comment", author="Alice")
        assert c.text == "New comment"
        assert c.author == "Alice"
        assert len(r1.comments) == 1
        assert r1.comments[0].id == 0
        assert p.comment_ids == [0]

        # Change the author, initials, and date of the comment
        c.author = "Bob"
        c.initials = "BB"
        c.date = "2024-01-01T00:00:00Z"
        assert c.author == "Bob"
        assert c.initials == "BB"
        assert c.date == "2024-01-01T00:00:00Z"

        # Added comment is also available in the paragraph's comments
        assert len(p.comments) == 1
        assert p.comments[0].text == "New comment"
        assert p.comments[0].id == 0

        # Add another comment, but to the paragraph
        c2 = p.add_comment("New comment 2", author="Charlie")
        assert c2.text == "New comment 2"
        assert len(p.comments) == 2
        assert p.comments[1].text == "New comment 2"
        assert p.comments[1].id == 1
