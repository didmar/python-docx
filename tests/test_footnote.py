"""Unit tests for the Footnote class."""

from docx import Document


class TestFootnote:
    def it_works(self):
        doc = Document()
        p = doc.add_paragraph("Hello world!")
        p.add_run(" Second run!")

        # Add a footnote to the paragraph
        f = p.add_footnote("New footnote")
        assert f.text == "New footnote"
        assert len(p.footnotes) == 1
        assert p.footnotes[0].id == 1
        assert p.footnote_ids == [1]

        # Check that the paragraph now has 3 runs: the first run, the second run, and the footnote
        assert len(p.runs) == 3
        assert p.runs[0].text == "Hello world!"
        assert p.runs[1].text == " Second run!"
        assert p.runs[2].footnote.id == 1

        # Change the text of the footnote
        f.text = "New footnote 2"
        assert f.text == "New footnote 2"

        # Add another footnote
        f2 = p.add_footnote("New footnote 3")
        assert f2.text == "New footnote 3"
        assert len(p.footnotes) == 2
        assert p.footnotes[1].id == 2
        assert p.footnote_ids == [1, 2]
